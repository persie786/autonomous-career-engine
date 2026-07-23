import json
import os
from google import genai
from google.genai import types
from docx import Document
from modules.ats_scorer import score_job_against_persona
from database.db_handler import (
    get_persona_for_company,
    save_generated_assets,
    log_activity,
    update_ats_score,
)

from database.db_handler import (
    get_persona_for_company,
    save_generated_assets,
    log_activity,
)
from modules.persona_builder import get_persona
from utils.logger import setup_logger
from utils.ai_router import generate_json

logger = setup_logger("cv_generator")

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(PROJECT_ROOT, "templates", "master_cv.docx")
OUTPUT_DIR = os.path.join(PROJECT_ROOT, "data", "generated_cvs")


GENERATION_PROMPT = """You are tailoring a candidate's CV and writing a cover letter for one \
specific job application.

Candidate persona (JSON):
{persona_json}

Job:
Title: {role}
Company: {company}
Description:
{description}

Return ONLY valid JSON in exactly this shape:
{{
  "tailored_summary": "a 2-3 sentence summary emphasizing fit for THIS role, based only on the persona's real background",
  "tailored_experience": "one line per bullet across all relevant experience entries, each prefixed with '- ', emphasizing points relevant to this role — never invent experience not present in the persona",
  "cover_letter": "a complete 3-4 paragraph cover letter for this specific role and company"
}}"""


def choose_persona_for_job(job: dict) -> tuple[str, dict]:
    """
    Tries, in order: the persona forced by the Consistency Guardrail (same
    company applied to before), then the persona for this job's search
    profile, then the manually-built 'default' persona. Only raises if none
    of the three exist — a missing or not-yet-created per-profile persona
    falls through instead of hard-failing.
    """
    tried = []

    forced_name = get_persona_for_company(job["company"])
    if forced_name:
        persona = get_persona(forced_name)
        if persona:
            logger.info(
                f"Consistency Guardrail: reusing persona '{forced_name}' for {job['company']}"
            )
            return forced_name, persona
        tried.append(f"company-forced '{forced_name}'")

    profile_name = job.get("search_profile")
    if profile_name:
        persona = get_persona(profile_name)
        if persona:
            return profile_name, persona
        tried.append(f"profile '{profile_name}'")

    persona = get_persona("default")
    if persona:
        return "default", persona
    tried.append("'default'")

    raise ValueError(
        f"No persona found for {job['company']} — tried {', '.join(tried)}. "
        "Build a persona in Settings first."
    )


def generate_application_assets(job: dict, persona: dict) -> dict:
    """Calls Gemini to produce tailored CV content and a cover letter for one job."""
    prompt = GENERATION_PROMPT.format(
        persona_json=json.dumps(persona, indent=2),
        role=job["role"],
        company=job["company"],
        description=job["job_description"][:4000],
    )
    raw_text, model_used = generate_json("", prompt, temperature=0.4)
    logger.info(f"Assets for job id={job['id']} generated via {model_used}")
    return json.loads(raw_text)


def _replace_tag(doc: Document, tag: str, replacement_text: str) -> bool:
    """Replaces the paragraph containing `tag` with one paragraph per line of
    replacement_text, preserving that paragraph's style."""
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            style = paragraph.style
            anchor = paragraph._p
            for line in replacement_text.split("\n") or [""]:
                new_para = doc.add_paragraph(line, style=style)
                anchor.addnext(new_para._p)
                anchor = new_para._p
            paragraph._p.getparent().remove(paragraph._p)
            return True
    return False


def build_docx(persona: dict, assets: dict, output_path: str):
    doc = Document(TEMPLATE_PATH)

    tags = {
        "{{FULL_NAME}}": persona.get("full_name", ""),
        "{{EMAIL}}": persona.get("email", ""),
        "{{PHONE}}": persona.get("phone", ""),
        "{{SUMMARY}}": assets["tailored_summary"],
        "{{SKILLS}}": ", ".join(persona.get("skills", [])),
        "{{EXPERIENCE}}": assets["tailored_experience"],
        "{{EDUCATION}}": "\n".join(
            f"{e.get('degree', '')} — {e.get('institution', '')} ({e.get('dates', '')})"
            for e in persona.get("education", [])
        ),
    }

    for tag, value in tags.items():
        if not _replace_tag(doc, tag, value):
            logger.warning(f"Tag {tag} not found in template — check master_cv.docx.")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)


def generate_for_job(job: dict) -> dict:
    """Full pipeline for one job — this is what the Live Asset Studio calls."""
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            "templates/master_cv.docx doesn't exist — add the {{TAGS}} listed in "
            "build_docx() to a real Word document at that path first."
        )

    persona_name, persona = choose_persona_for_job(job)
    ats_result = score_job_against_persona(job["job_description"], persona)
    update_ats_score(
        job["id"], ats_result["total"], ats_result["matched"], ats_result["missing"]
    )
    assets = generate_application_assets(job, persona)

    safe_company = "".join(
        c for c in job["company"] if c.isalnum() or c in " -_"
    ).strip()
    output_path = os.path.join(OUTPUT_DIR, f"CV_{safe_company}_{job['id']}.docx")
    build_docx(persona, assets, output_path)

    save_generated_assets(
        job_id=job["id"],
        persona_name=persona_name,
        cv_text=f"{assets['tailored_summary']}\n\n{assets['tailored_experience']}",
        cover_letter_text=assets["cover_letter"],
        docx_path=output_path,
    )

    logger.info(
        f"Generated assets for job id={job['id']} ({job['role']} at {job['company']})"
    )
    return {"persona_used": persona_name, "docx_path": output_path, **assets}


if __name__ == "__main__":
    from database.db_handler import get_jobs

    pending = get_jobs(status="Pending")
    if not pending:
        print("No 'Pending' jobs to generate assets for.")
    else:
        job = pending[0]
        print(f"Generating assets for: {job['role']} at {job['company']}")
        result = generate_for_job(job)
        print(f"\nPersona used: {result['persona_used']}")
        print(f"Docx saved to: {result['docx_path']}")
        print(f"\n--- Summary ---\n{result['tailored_summary']}")
        print(f"\n--- Cover Letter ---\n{result['cover_letter']}")
